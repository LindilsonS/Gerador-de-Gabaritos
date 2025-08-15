import os
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import filedialog, messagebox, Listbox, Scrollbar, Frame, Radiobutton, StringVar
import re
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle

def sanitizar_nome(nome):
    nome_str = str(nome)
    caracteres_invalidos = r'[<>:"/\\|?*]'
    return re.sub(caracteres_invalidos, '_', nome_str).strip()

def get_unique_teachers(df):
    """Get unique staff members from dataframe sorted alphabetically"""
    # Check which format we're dealing with
    if 'CPF DO PROFESSOR' in df.columns:
        funcionarios = df.drop_duplicates(subset=['NOME DO PROFESSOR', 'CPF DO PROFESSOR'])[['NOME DO PROFESSOR', 'CPF DO PROFESSOR']]
        funcionarios = funcionarios.rename(columns={'NOME DO PROFESSOR': 'NOME DO FUNCIONÁRIO'})
    else:
        # For old format
        funcionarios = df.drop_duplicates(subset=['PROFESSOR REGENTE'])[['PROFESSOR REGENTE']]
        funcionarios = funcionarios.rename(columns={'PROFESSOR REGENTE': 'NOME DO FUNCIONÁRIO'})
    return funcionarios.sort_values('NOME DO FUNCIONÁRIO')

def detect_csv_format(df):
    """Detect which CSV format we're dealing with"""
    if {'NOME DO PROFESSOR', 'CPF DO PROFESSOR', 'TURNO'}.issubset(df.columns):
        return "professor_format"
    elif {'PROFESSOR REGENTE', 'NOME DO ALUNO', 'ETAPA DE ENSINO'}.issubset(df.columns):
        return "aluno_format"
    else:
        raise ValueError("Formato de CSV não reconhecido")

# --- FUNÇÃO DE SUBSTITUIÇÃO COM A CORREÇÃO DA ORDEM DE RECONHECIMENTO ---
def substituir_variaveis_em_tudo(doc, dados):
    """
    Substitui placeholders em todo o documento, incluindo parágrafos, tabelas e caixas de texto.
    A correção crucial aqui é ordenar as chaves por comprimento para evitar substituições parciais.
    """
    # Helper function para fazer a substituição de forma segura
    def substituir_texto(texto_original, dados_substituicao):
        # ORDENA AS CHAVES (placeholders) POR COMPRIMENTO, DA MAIOR PARA A MENOR.
        # Isso garante que '$VARIÁVEL NOME DO ALUNO 2' seja processado ANTES de '$VARIÁVEL NOME DO ALUNO'.
        for key in sorted(dados_substituicao.keys(), key=len, reverse=True):
            value = dados_substituicao[key]
            texto_original = texto_original.replace(key, str(value))
        return texto_original

    # 1. Substituição em parágrafos e tabelas do corpo principal
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    
    for p in all_paragraphs:
        full_text = "".join(run.text for run in p.runs)
        if '$' in full_text:
            novo_texto = substituir_texto(full_text, dados)
            # Apenas modifica o parágrafo se houver mudança, para preservar formatação
            if novo_texto != full_text:
                for run in p.runs:
                    run.text = ''
                p.add_run(novo_texto)

    # 2. Substituição DENTRO DE CAIXAS DE TEXTO (acessando o XML)
    W_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    TAG_TEXTBOX = f"{W_NAMESPACE}txbxContent"
    TAG_PARAGRAPH = f"{W_NAMESPACE}p"
    TAG_RUN = f"{W_NAMESPACE}r"
    TAG_TEXT = f"{W_NAMESPACE}t"

    for txbx in doc.element.body.iter(TAG_TEXTBOX):
        for p_element in txbx.iter(TAG_PARAGRAPH):
            full_text = ""
            all_t_elements = []
            for r_element in p_element.iter(TAG_RUN):
                for t_element in r_element.iter(TAG_TEXT):
                    all_t_elements.append(t_element)
                    if t_element.text:
                        full_text += t_element.text

            if '$' in full_text:
                novo_texto = substituir_texto(full_text, dados)
                if novo_texto != full_text:
                    for i, t_element in enumerate(all_t_elements):
                        if i == 0:
                            t_element.text = novo_texto
                        else:
                            t_element.text = ""
    
    return doc

def criar_lista_presenca(escola, nome_grupo, items, diretorio_saida, titulo_lista="Lista de Presença", cores=None, data_lista=None, is_teacher_list=False):
    items = sorted(items) if not is_teacher_list else items
    caminho_arquivo = os.path.join(diretorio_saida, f"lista_presenca_{nome_grupo}.pdf")
    c = canvas.Canvas(caminho_arquivo, pagesize=A4)
    largura, altura = A4
    margem_esquerda, margem_superior, margem_inferior = 40, altura - 40, 40
    # Ajustando a distribuição das colunas
    if is_teacher_list:
        col_widths = [28, 250, 228.35]  # Número, Nome do Funcionário, Assinatura
    else:
        col_widths = [28, 228.35, 250]  # Mantém o padrão original para alunos
    row_height, font_size = 30, 12
    max_rows_per_page = 20

    def desenhar_cabecalho_rodape(ultima_pagina=False):
        c.setFont("Helvetica-Bold", font_size + 8)
        c.setFillColor(colors.HexColor(cores["titulo"]))
        c.drawCentredString(largura / 2, margem_superior, titulo_lista)
        c.setFont("Helvetica-Bold", font_size)
        c.setFillColor(colors.HexColor(cores["cabecalho"]))
        c.drawCentredString(largura / 2, margem_superior - 20, f"Escola: {escola}")
        c.setFont("Helvetica", font_size)
        c.setFillColor(colors.black)
        texto_data = f"Data: {data_lista}" if data_lista else "Data: ____________"
        c.drawString(margem_esquerda, margem_superior - 40, texto_data)
        c.setStrokeColor(colors.HexColor(cores["linha"]))
        c.setLineWidth(1)
        c.line(margem_esquerda, margem_superior - 50, largura - 40, margem_superior - 50)
        
        if ultima_pagina:
            c.drawString(margem_esquerda, margem_inferior + 40, f"Total de {'funcionários' if is_teacher_list else 'alunos'}: {len(items)}")
            c.drawString(margem_esquerda, margem_inferior + 20, f"Total de {'funcionários' if is_teacher_list else 'alunos'} presentes: ________")
            c.drawCentredString(largura / 2, margem_inferior, "SECRETARIA MUNICIPAL DE EDUCAÇÃO")

    total_paginas = (len(items) + max_rows_per_page - 1) // max_rows_per_page
    
    for page, start in enumerate(range(0, len(items), max_rows_per_page)):
        if page > 0: c.showPage()
        e_ultima_pagina = (page == total_paginas - 1)
        desenhar_cabecalho_rodape(ultima_pagina=e_ultima_pagina)
        
        header = ["Nº", "Nome do Funcionário" if is_teacher_list else "Nome do Aluno", "Assinatura"]
        if is_teacher_list:
            dados_tabela = [header] + [[str(i + 1 + start), item['NOME DO FUNCIONÁRIO'], ""] for i, item in enumerate(items[start:start + max_rows_per_page].to_dict('records'))]
        else:
            dados_tabela = [header] + [[str(i + 1 + start), item, ""] for i, item in enumerate(items[start:start + max_rows_per_page])]
            
        table = Table(dados_tabela, colWidths=col_widths, rowHeights=row_height)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(cores["tabela_header"])),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor(cores["titulo"])),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor("#ECF0F1")),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        altura_tabela = margem_superior - 70
        table.wrapOn(c, largura, altura)
        table.drawOn(c, margem_esquerda, altura_tabela - len(dados_tabela) * row_height)
    c.save()
    print(f"Lista de presença salva: {caminho_arquivo}")

def criar_gabaritos(csv_path, modelo_path, output_dir, config, etapas_selecionadas, escolas_selecionadas):
    try:
        df = pd.read_csv(csv_path, sep=';', encoding='utf-8', dtype=str)
        csv_format = detect_csv_format(df)
        
        if csv_format == "professor_format":
            df = df.rename(columns={
                'ETAPA': 'ETAPA DE ENSINO',
                'NOME DA ESCOLA': 'ESCOLA'
            })
        
        if config.is_teacher_list:
            if csv_format == "professor_format":
                colunas_necessarias = {'ESCOLA', 'NOME DO PROFESSOR', 'CPF DO PROFESSOR'}
            else:
                colunas_necessarias = {'ESCOLA', 'PROFESSOR REGENTE'}
        else:
            if csv_format == "professor_format":
                raise ValueError("Este arquivo CSV contém apenas dados de professores")
            colunas_necessarias = {'ESCOLA', 'TURMA', 'NOME DO ALUNO', 'PROFESSOR REGENTE', 'ETAPA DE ENSINO'}
            
        if not colunas_necessarias.issubset(df.columns):
            raise ValueError(f"O arquivo CSV deve conter as colunas: {', '.join(colunas_necessarias)}")

        if escolas_selecionadas is None:
            escolas_selecionadas = df['ESCOLA'].dropna().unique()

        df_filtrado = df[df['ESCOLA'].fillna('').isin(escolas_selecionadas)].copy()
        
        if df_filtrado.empty:
            return True, "Nenhum registro encontrado para as escolas selecionadas."

        for escola in escolas_selecionadas:
            escola_df = df_filtrado[df_filtrado['ESCOLA'] == escola]
            escola_sanitizada = sanitizar_nome(escola)
            base_dir = os.path.join(output_dir, escola_sanitizada)
            os.makedirs(base_dir, exist_ok=True)

            if config.is_teacher_list:
                funcionarios = get_unique_teachers(escola_df)
                criar_lista_presenca(
                    escola, 
                    "funcionarios",
                    funcionarios,
                    base_dir,
                    config.titulo_lista,
                    config.cores,
                    config.data_lista,
                    is_teacher_list=True
                )
            else:
                alunos = escola_df['NOME DO ALUNO'].tolist()
                professor_regente = escola_df['PROFESSOR REGENTE'].iloc[0] if not escola_df['PROFESSOR REGENTE'].empty else "Não especificado"
                
                if config.gerar_lista_presenca or config.apenas_lista_presenca:
                    criar_lista_presenca(escola, turma_sanitizada, alunos, base_dir, 
                                       config.titulo_lista, config.cores, config.data_lista)
                
                # Só gera os gabaritos se não estiver no modo "apenas lista de presença"
                if not config.apenas_lista_presenca:
                    if config.process_mode == "um_aluno":
                        for aluno in alunos:
                            doc = Document(modelo_path)
                            dados_aluno = {
                                '$VARIÁVEL ESCOLA': escola,
                                '$VARIÁVEL TURMA': turma,
                                '$VARIÁVEL PROFESSOR REGENTE': professor_regente,
                                '$VARIÁVEL NOME DO ALUNO': aluno,
                                '$VARIÁVEL NOME DO ALUNO 2': ''
                            }
                            substituir_variaveis_em_tudo(doc, dados_aluno)
                            nome_arquivo = f"{sanitizar_nome(aluno)}_gabarito.docx"
                            doc.save(os.path.join(base_dir, nome_arquivo))
                            print(f"Arquivo salvo (1 aluno): {nome_arquivo}")

                    else:
                        for i in range(0, len(alunos), 2):
                            aluno1 = alunos[i]
                            aluno2 = alunos[i + 1] if (i + 1) < len(alunos) else None
                            doc = Document(modelo_path)
                            dados_alunos = {
                                '$VARIÁVEL ESCOLA': escola,
                                '$VARIÁVEL TURMA': turma,
                                '$VARIÁVEL PROFESSOR REGENTE': professor_regente,
                                '$VARIÁVEL NOME DO ALUNO': aluno1,
                                '$VARIÁVEL NOME DO ALUNO 2': aluno2 if aluno2 else ''
                            }
                            substituir_variaveis_em_tudo(doc, dados_alunos)
                            aluno1_sanitizado = sanitizar_nome(aluno1)
                            if aluno2:
                                aluno2_sanitizado = sanitizar_nome(aluno2)
                                nome_arquivo = f"{aluno1_sanitizado}_e_{aluno2_sanitizado}_gabarito.docx"
                            else:
                                nome_arquivo = f"{aluno1_sanitizado}_gabarito.docx"
                            doc.save(os.path.join(base_dir, nome_arquivo))
                            print(f"Arquivo salvo (2 alunos): {nome_arquivo}")

        return True, "Documentos gerados com sucesso!"
    except Exception as e:
        import traceback
        traceback.print_exc()
        return False, str(e)

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Gerador de Gabaritos e Lista de Presença")
        self.root.geometry("600x780")  
        self.root.minsize(580, 860) 
        
        # Inicialização de todas as variáveis no início
        self.csv_path = None
        self.modelo_path = None
        self.output_dir = None
        self.process_mode = StringVar(value="dois_alunos")
        self.gerar_lista_presenca = tk.BooleanVar(value=True)
        self.apenas_lista_presenca = tk.BooleanVar(value=False)  # Nova variável
        self.titulo_lista = tk.StringVar(value="Lista de Presença")
        self.data_lista = tk.StringVar(value="")  # Nova variável para data
        self.paleta_selecionada = StringVar(value="Verde Suave")
        self.lista_tipo = StringVar(value="alunos")  # Add this after other initializations
        
        # Paletas de cores pasteis pré-definidas
        self.paletas_cores = {
            "Verde Suave": {
                "titulo": "#2C3E50",
                "cabecalho": "#34495E",
                "linha": "#3ddb65",
                "tabela_header": "#98FB98"
            },
            "Rosa Delicado": {
                "titulo": "#4A4A4A",
                "cabecalho": "#5D4E6D",
                "linha": "#FFB6C1",
                "tabela_header": "#FFC0CB"
            },
            "Azul Sereno": {
                "titulo": "#2C3E50",
                "cabecalho": "#34495E",
                "linha": "#87CEEB",
                "tabela_header": "#ADD8E6"
            },
            "Lilás Suave": {
                "titulo": "#4A4A4A",
                "cabecalho": "#5D4E6D",
                "linha": "#DDA0DD",
                "tabela_header": "#E6E6FA"
            },
            "Marrom Café": {
                "titulo": "#3E2723",
                "cabecalho": "#4E342E",
                "linha": "#8D6E63",
                "tabela_header": "#D7CCC8"
            },
            "Cinza Elegante": {
                "titulo": "#263238",
                "cabecalho": "#37474F",
                "linha": "#78909C",
                "tabela_header": "#CFD8DC"
            },
            "Verde Menta": {
                "titulo": "#004D40",
                "cabecalho": "#00695C",
                "linha": "#4DB6AC",
                "tabela_header": "#B2DFDB"
            },
            "Roxo Real": {
                "titulo": "#311B92",
                "cabecalho": "#4527A0",
                "linha": "#7E57C2",
                "tabela_header": "#D1C4E9"
            },
            "Laranja Solar": {
                "titulo": "#E65100",
                "cabecalho": "#EF6C00",
                "linha": "#FFB74D",
                "tabela_header": "#FFE0B2"
            },
            "Azul Corporativo": {
                "titulo": "#0D47A1",
                "cabecalho": "#1565C0",
                "linha": "#42A5F5",
                "tabela_header": "#BBDEFB"
            }
        }
        
        tk.Label(root, text="Gerador de Gabaritos", font=("Arial", 14, "bold")).pack(pady=10)
        
        # Adicionar campo para título da lista e data
        titulo_frame = Frame(root)
        titulo_frame.pack(pady=5, padx=10, fill=tk.X)
        tk.Label(titulo_frame, text="Título da Lista:").pack(side=tk.LEFT)
        tk.Entry(titulo_frame, textvariable=self.titulo_lista, width=40).pack(side=tk.LEFT, padx=5)
        
        data_frame = Frame(root)
        data_frame.pack(pady=5, padx=10, fill=tk.X)
        tk.Label(data_frame, text="Data da Lista:").pack(side=tk.LEFT)
        tk.Entry(data_frame, textvariable=self.data_lista, width=40).pack(side=tk.LEFT, padx=5)
        tk.Label(data_frame, text="(Opcional - deixe em branco para preenchimento manual)").pack(side=tk.LEFT, padx=5)
        
        frame_botoes = Frame(root)
        frame_botoes.pack(pady=5, padx=10, fill=tk.X)
        
        tk.Button(frame_botoes, text="1. Selecionar CSV dos Alunos", command=self.selecionar_csv).pack(fill=tk.X)
        self.csv_label = tk.Label(frame_botoes, text="Nenhum arquivo selecionado", fg="blue")
        self.csv_label.pack()
        
        tk.Button(frame_botoes, text="2. Selecionar Modelo Word (.docx)", command=self.selecionar_modelo).pack(fill=tk.X, pady=(10,0))
        self.modelo_label = tk.Label(frame_botoes, text="Nenhum arquivo selecionado", fg="blue")
        self.modelo_label.pack()
        
        tk.Button(frame_botoes, text="3. Selecionar Pasta de Saída", command=self.selecionar_output).pack(fill=tk.X, pady=(10,0))
        self.output_label = tk.Label(frame_botoes, text="Nenhuma pasta selecionada", fg="blue")
        self.output_label.pack()

        # Após os botões de seleção de arquivos, adicionar seletor de escolas
        tk.Label(root, text="4. Selecione as Escolas:", font=("Arial", 10, "bold")).pack(pady=(15, 0))
        
        escolas_frame = Frame(root)
        escolas_frame.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)
        escolas_scrollbar = Scrollbar(escolas_frame, orient=tk.VERTICAL)
        self.escolas_listbox = Listbox(
            escolas_frame,
            selectmode=tk.MULTIPLE,
            yscrollcommand=escolas_scrollbar.set,
            exportselection=False,
            height=6
        )
        escolas_scrollbar.config(command=self.escolas_listbox.yview)
        escolas_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.escolas_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        tk.Label(root, text="5. Selecione as Etapas de Ensino:", font=("Arial", 10, "bold")).pack(pady=(15, 0))
        
        etapas_frame = Frame(root)
        etapas_frame.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)
        scrollbar = Scrollbar(etapas_frame, orient=tk.VERTICAL)
        self.etapas_listbox = Listbox(
            etapas_frame, 
            selectmode=tk.MULTIPLE, 
            yscrollcommand=scrollbar.set, 
            exportselection=False,
            height=8  
        )
        scrollbar.config(command=self.etapas_listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.etapas_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        opcoes_frame = Frame(root)
        opcoes_frame.pack(pady=10)
        tk.Label(opcoes_frame, text="Modo de Geração:", font=("Arial", 10, "bold")).pack(side=tk.LEFT, padx=(0, 10))
        Radiobutton(opcoes_frame, text="Um aluno por folha", variable=self.process_mode, value="um_aluno").pack(side=tk.LEFT)
        Radiobutton(opcoes_frame, text="Dois alunos por folha", variable=self.process_mode, value="dois_alunos").pack(side=tk.LEFT)

        # Adicionar os checkboxes em um frame separado
        check_frame = Frame(root)
        check_frame.pack(pady=5)
        tk.Checkbutton(check_frame, text="Gerar também a Lista de Presença em PDF", 
                      variable=self.gerar_lista_presenca).pack(side=tk.LEFT, padx=5)
        tk.Checkbutton(check_frame, text="Gerar apenas Lista de Presença", 
                      variable=self.apenas_lista_presenca).pack(side=tk.LEFT, padx=5)

        botao_gerar = tk.Button(
            root, 
            text="GERAR DOCUMENTOS", 
            command=self.gerar,
            font=("Arial", 12, "bold"),
            bg="#2ecc71",  
            fg="white",
            relief=tk.RAISED,
            bd=2,
            padx=20,
            pady=10,
            cursor="hand2" 
        )
        botao_gerar.pack(pady=20)
        
        # Adiciona efeito hover
        def on_enter(e):
            botao_gerar['bg'] = '#27ae60'  # Verde um pouco mais escuro ao passar o mouse
            
        def on_leave(e):
            botao_gerar['bg'] = '#2ecc71'  # Volta ao verde original
            
        botao_gerar.bind("<Enter>", on_enter)
        botao_gerar.bind("<Leave>", on_leave)
        
        # Adicionar seleção de cores
        cores_frame = Frame(root)
        cores_frame.pack(pady=5, padx=10, fill=tk.X)
        tk.Label(cores_frame, text="Paleta de Cores:", font=("Arial", 10, "bold")).pack(side=tk.LEFT)
        
        for nome_paleta in self.paletas_cores.keys():
            rb = Radiobutton(
                cores_frame,
                text=nome_paleta,
                variable=self.paleta_selecionada,
                value=nome_paleta,
                command=self.atualizar_preview_cores
            )
            rb.pack(side=tk.LEFT, padx=5)
        
        # Preview das cores
        self.preview_frame = Frame(root)
        self.preview_frame.pack(pady=5, padx=10)
        self.atualizar_preview_cores()
    
        tipo_lista_frame = Frame(root)
        tipo_lista_frame.pack(pady=5)
        tk.Label(tipo_lista_frame, text="Tipo de Lista:").pack(side=tk.LEFT)
        Radiobutton(tipo_lista_frame, text="Lista de Alunos", variable=self.lista_tipo, 
                   value="alunos", command=self.atualizar_modo_lista).pack(side=tk.LEFT)
        Radiobutton(tipo_lista_frame, text="Lista de Funcionários", variable=self.lista_tipo,
                   value="professores", command=self.atualizar_modo_lista).pack(side=tk.LEFT)

    def atualizar_modo_lista(self):
        is_teacher_list = self.lista_tipo.get() == "professores"
        if is_teacher_list:
            self.process_mode.set("um_aluno")  # Force single mode for teacher lists
            self.apenas_lista_presenca.set(True)  # Force presence list only
        
    def atualizar_preview_cores(self):
        for widget in self.preview_frame.winfo_children():
            widget.destroy()
            
        paleta = self.paletas_cores[self.paleta_selecionada.get()]
        for nome, cor in paleta.items():
            preview = Frame(self.preview_frame, bg=cor, width=30, height=20)
            preview.pack(side=tk.LEFT, padx=2)
            preview.pack_propagate(False)
            tk.Label(preview, text="", bg=cor).pack(expand=True)
    
    def selecionar_csv(self):
        self.csv_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv")])
        if self.csv_path:
            self.csv_label.config(text=os.path.basename(self.csv_path))
            self.popular_lista_etapas()
    
    def popular_lista_etapas(self):
        try:
            self.etapas_listbox.delete(0, tk.END)
            self.escolas_listbox.delete(0, tk.END)
            
            df = pd.read_csv(self.csv_path, sep=';', encoding='utf-8', dtype=str)
            csv_format = detect_csv_format(df)
            
            # Normalize column names based on format
            if csv_format == "professor_format":
                df = df.rename(columns={
                    'ETAPA': 'ETAPA DE ENSINO',
                    'NOME DA ESCOLA': 'ESCOLA'
                })
            
            coluna_etapa = 'ETAPA DE ENSINO'
            coluna_escola = 'ESCOLA'
            
            if coluna_escola not in df.columns:
                messagebox.showerror("Erro de CSV", 
                    "O arquivo CSV precisa ter a coluna 'ESCOLA' ou 'NOME DA ESCOLA'.")
                return
                
            # Popular lista de escolas
            escolas_unicas = sorted(df[coluna_escola].dropna().unique())
            for escola in escolas_unicas:
                self.escolas_listbox.insert(tk.END, escola)
            
            # Popular lista de etapas se existirem
            if coluna_etapa in df.columns:
                etapas_unicas = sorted(df[coluna_etapa].dropna().unique())
                for etapa in etapas_unicas:
                    self.etapas_listbox.insert(tk.END, etapa)
            
        except Exception as e:
            messagebox.showerror("Erro ao Ler CSV", f"Não foi possível processar o arquivo CSV: {str(e)}")
            self.csv_path = None
            self.csv_label.config(text="Nenhum CSV selecionado")

    def selecionar_modelo(self):
        self.modelo_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if self.modelo_path:
            self.modelo_label.config(text=os.path.basename(self.modelo_path))
    
    def selecionar_output(self):
        self.output_dir = filedialog.askdirectory()
        if self.output_dir:
            self.output_label.config(text=self.output_dir)
    
    def gerar(self):
        if not all([self.csv_path, 
                    (self.modelo_path or self.apenas_lista_presenca.get()), 
                    self.output_dir]):
            messagebox.showerror("Erro", 
                "Por favor, selecione o arquivo CSV, a pasta de saída " + 
                ("e o modelo Word!" if not self.apenas_lista_presenca.get() else "!"))
            return
        
        # Se nada estiver selecionado, usa todas as etapas e escolas
        etapas_selecionadas = [self.etapas_listbox.get(i) for i in self.etapas_listbox.curselection()] \
            if self.etapas_listbox.curselection() else None
        escolas_selecionadas = [self.escolas_listbox.get(i) for i in self.escolas_listbox.curselection()] \
            if self.escolas_listbox.curselection() else None
        
        try:
            config = Configuracao(
                process_mode=self.process_mode.get(),
                gerar_lista_presenca=self.gerar_lista_presenca.get(),
                apenas_lista_presenca=self.apenas_lista_presenca.get(),
                titulo_lista=self.titulo_lista.get(),
                data_lista=self.data_lista.get(),
                cores=self.paletas_cores[self.paleta_selecionada.get()],
                is_teacher_list=self.lista_tipo.get() == "professores"
            )
            sucesso, mensagem = criar_gabaritos(self.csv_path, self.modelo_path, self.output_dir, 
                                              config, etapas_selecionadas, escolas_selecionadas)
            
            if sucesso:
                messagebox.showinfo("Sucesso", mensagem)
            else:
                messagebox.showerror("Erro", f"Falha ao gerar os documentos.\n\nDetalhes: {mensagem}")
        except Exception as e:
            messagebox.showerror("Erro Inesperado", f"Ocorreu um erro inesperado: {str(e)}")

class Configuracao:
    def __init__(self, process_mode="dois_alunos", gerar_lista_presenca=True,
                 apenas_lista_presenca=False, titulo_lista="Lista de Presença",
                 data_lista="", cores=None, is_teacher_list=False):
        self.process_mode = process_mode
        self.gerar_lista_presenca = gerar_lista_presenca
        self.apenas_lista_presenca = apenas_lista_presenca
        self.titulo_lista = titulo_lista
        self.data_lista = data_lista 
        self.cores = cores or {
            "titulo": "#2C3E50",
            "cabecalho": "#34495E",
            "linha": "#3ddb65",
            "tabela_header": "#98FB98"
        }
        self.is_teacher_list = is_teacher_list

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()